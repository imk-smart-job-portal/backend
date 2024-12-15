import json
import sqlite3
# from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
# from unsloth import FastLanguageModel 
from pydantic import BaseModel
from passlib.context import CryptContext
import jwt as pyjwt 
from datetime import datetime, timedelta
from fastapi import File, UploadFile
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO

# Konfigurasi logging
import logging

# Setup logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Test logging output
logger.debug("This is a debug message")


# Middleware untuk mengizinkan akses dari berbagai sumber (CORS)
origins = ["*"]
from dotenv import load_dotenv
from auth import hash_password_hashlib, create_jwt, verify_jwt
from database import get_db_connection, initialize_database
from jobs import router as jobs_router
from genai import generate_required_skills

load_dotenv()

app = FastAPI()

# Middleware CORS
origins = ["*"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

max_seq_length = 2048 
dtype = None
load_in_4bit = True

SECRET_KEY = "your-secret-key"  
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30

DATABASE = "job_portal.db"
conn = sqlite3.connect(DATABASE, check_same_thread=False)
conn.row_factory = sqlite3.Row

initialize_database()

#Helper Ekstrasi
def extract_text_from_pdf(file_content: bytes) -> str:
    """Ekstrak teks dari file PDF."""
    pdf_reader = PdfReader(BytesIO(file_content))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text.strip()

def extract_text_from_docx(file_content: bytes) -> str:
    """Ekstrak teks dari file DOCX."""
    doc = Document(BytesIO(file_content))
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text.strip()

app = FastAPI()

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Models
class RegisterRequest(BaseModel):
    name: str
    email: str
    phone_number: str
    password: str

class LoginRequest(BaseModel):
    email: str
    password: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str

# Helper Functions
def create_token(data: dict, expires_delta: timedelta = timedelta(hours=1)):
    to_encode = data.copy()
    expire = datetime.utcnow() + expires_delta
    to_encode.update({"exp": expire})
    return pyjwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def hash_password(password: str):
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str):
    logger.debug(f"Verifying password: {plain_password} with hash: {hashed_password}")
    return pwd_context.verify(plain_password, hashed_password)


# Endpoints register
@app.post("/register", response_model=TokenResponse)
def register_user(user: RegisterRequest):
    cursor = conn.cursor()

    # Cek apakah role dengan name = user sudah ada
    cursor.execute("SELECT * FROM roles WHERE name = ?", ('user',))
    existing_role_user = cursor.fetchone()

    if not existing_role_user:
        cursor.execute("""
            INSERT INTO roles (id, name)
            VALUES (?, ?)
        """, (1, 'user'))
        conn.commit()

    cursor.execute("SELECT * FROM applicants WHERE email = ?", (user.email,))
    if cursor.fetchone():
        raise HTTPException(status_code=400, detail="Email already registered")
    
    hashed_password = hash_password_hashlib(user.password)
    cursor.execute("""
        INSERT INTO applicants (role_id, name, email, phone_number, password)
        VALUES (?, ?, ?, ?, ?)
    """, (1, user.name, user.email, user.phone_number, hashed_password))
    conn.commit()
    
    token = create_token({"sub": user.email})
    return {"access_token": token, "token_type": "bearer"}
#endpoint login
@app.post("/login", response_model=TokenResponse)
def login_user(credentials: LoginRequest):
    logger.debug(f"Login attempt with email: {credentials.email}")
    cursor = conn.cursor()
    hashed_password = hash_password_hashlib(credentials.password)
    try:
        cursor.execute("SELECT * FROM applicants WHERE email = ? AND password = ?", (credentials.email, hashed_password))
        user = cursor.fetchone()

        # Check if user exists and password matches
        if not user:
            logger.debug(f"User with email {credentials.email} not found.")
            raise HTTPException(status_code=400, detail="Invalid credentials")

        # Create token if login is successful
        token = create_token({"sub": user["email"], "applicant_id": user["id"]})
        logger.debug(f"Login successful for user: {credentials.email}")

        return {"access_token": token, "token_type": "bearer"}
    
    except Exception as e:
        logger.debug(f"Error during login: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

#endpoint upload resume
@app.post("/upload-resume")
async def upload_resume(token: dict = Depends(verify_jwt), file: UploadFile = File(...)):
    applicant_id = token['applicant_id']
    
    # Validasi format file
    if file.content_type == "application/pdf":
        # Baca teks dari file PDF
        reader = PdfReader(file.file)
        text = " ".join([page.extract_text() for page in reader.pages])
    elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        # Baca teks dari file DOCX
        doc = Document(file.file)
        text = " ".join([p.text for p in doc.paragraphs])
    else:
        # Jika format file tidak valid
        raise HTTPException(status_code=400, detail="Unsupported file format")

    required_skills = generate_required_skills(text)

    print(required_skills)
    # Simpan string JSON ke database
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO resumes (applicant_id, files, required_skills) VALUES (?, ?, ?)",
        (applicant_id, text, required_skills)
    )
    conn.commit()

    return {"message": "Resume uploaded and processed successfully"}

#Endpoint cek resume
@app.get("/resume/{applicant_id}")
def get_resume(applicant_id: int):
    cursor = conn.cursor()
    cursor.execute("SELECT files FROM resumes WHERE applicant_id = ?", (applicant_id,))
    result = cursor.fetchone()
    if result:
        # Ubah string JSON kembali ke objek JSON
        resume_json = json.loads(result["files"])
        return resume_json
    raise HTTPException(status_code=404, detail="Resume not found")

@app.put("/resume/{resume_id}")
async def update_resume(resume_id: int, file: UploadFile = File(...)):
    # Cek apakah resume dengan resume_id ada
    cursor = conn.cursor()
    cursor.execute("SELECT applicant_id FROM resumes WHERE id = ?", (resume_id,))
    result = cursor.fetchone()

    if not result:
        raise HTTPException(status_code=404, detail="Resume not found")

    applicant_id = result[0]  # Ambil applicant_id dari database

    # Validasi format file
    if file.content_type == "application/pdf":
        # Baca teks dari file PDF
        reader = PdfReader(file.file)
        text = " ".join([page.extract_text() for page in reader.pages])
    elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        # Baca teks dari file DOCX
        doc = Document(file.file)
        text = " ".join([p.text for p in doc.paragraphs])
    else:
        # Jika format file tidak valid
        raise HTTPException(status_code=400, detail="Unsupported file format")

    # Konversi teks mentah ke JSON terstruktur (contoh sederhana)
    # parsed_json = {
    #     "applicant_id": applicant_id,
    #     "raw_text": text,
    #     "name": None,  # Ekstraksi nama (gunakan regex atau NLP di masa depan)
    #     "email": None,  # Ekstraksi email
    #     "skills": [],   # Ekstraksi skill
    #     "education": None,  # Ekstraksi pendidikan
    #     "experience": None  # Ekstraksi pengalaman kerja
    # }

    # Ubah JSON menjadi string sebelum disimpan
    # json_string = json.dumps(parsed_json)

    # Update resume di database
    cursor.execute(
        "UPDATE resumes SET files = ? WHERE id = ?",
        (text, resume_id)
    )
    conn.commit()

    return {"message": "Resume updated successfully"}

#endpoint delete resume
@app.delete("/resume/{resume_id}")
def delete_resume(resume_id: int):
    # Cek apakah resume ada di database
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM resumes WHERE id = ?", (resume_id,))
    resume = cursor.fetchone()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Hapus resume dari database
    cursor.execute("DELETE FROM resumes WHERE id = ?", (resume_id,))
    conn.commit()

    return {"message": "Resume deleted successfully"}

# Endpoint login dan mendapatkan JWT
@app.post('/companies/login')
def login_company(email: str, password: str):
    hashed_password = hash_password_hashlib(password)
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM companies WHERE email = ? AND password = ?', (email, hashed_password))
        company = cursor.fetchone()
        if company:
            token = create_jwt(company[0], company[2])  # company[0] = company_id, company[2] = company_name
            return {"message": "Login successful", "access_token": token}
        else:
            raise HTTPException(status_code=401, detail="Invalid email or password")

# Endpoint untuk mendapatkan data perusahaan
@app.get("/companies/me")
def get_company_data(token: str = Depends(verify_jwt)):
    company_id = token['company_id']  # Mengambil company_id dari token
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM companies WHERE id = ?", (company_id,))
        company = cursor.fetchone()
        if company:
            return {
                "id": company[0],              # id
                "role_id": company[1],         # role_id
                "name": company[2],            # name
                "email": company[3],           # email
                "phone_number": company[4],    # phone_number
            }
        else:
            raise HTTPException(status_code=404, detail="Company not found")

# Endpoint untuk mendaftar perusahaan
@app.post('/companies/register')
def register_company(email: str, password: str, company_name: str, phone_number: str):
    hashed_password = hash_password_hashlib(password)
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO companies (email, password, name, phone_number, role_id)
                VALUES (?, ?, ?, ?, ?)
            ''', (email, hashed_password, company_name, phone_number, 2))
            conn.commit()
        return {"message": "Company registered successfully"}
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="Email already exists")

# Endpoint untuk mendapatkan daftar pekerjaan
@app.get('/job-vacancies')
def job_vacancies():
    # Inisialisasi database
    initialize_database()
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM job_vacancies JOIN companies ON job_vacancies.company_id = companies.id')
        data = cursor.fetchall()
        outs = []
        for d in data:
            out = {
                'id': d[0],
                'company_id': d[1],
                'name': d[2],
                'description': d[3],
                'requirement': d[4],
                'created_at': d[5],
                'updated_at': d[6],
                'company': {
                    'id': d[7],
                    'user_id': d[8],
                    'created_at': d[11],
                    'updated_at': d[12],
                },
            }
            outs.append(out)
        return outs
    
app.include_router(jobs_router)


def get_db():
    conn = sqlite3.connect("job_portal.db")
    conn.row_factory = sqlite3.Row
    try:
        yield conn
    finally:
        conn.close()

# Endpoint untuk mendapatkan daftar tabel dalam database SQLite
@app.get('/tables')
def list_tables(db: sqlite3.Connection = Depends(get_db)):
    q = "SELECT name FROM sqlite_master WHERE type='table';"
    tables = conn.execute(q).fetchall()
    return [table['name'] for table in tables]

@app.get('/tables/companies')
def inspect_table_companies(db: sqlite3.Connection = Depends(get_db)):
    q = "PRAGMA table_info(companies);"
    columns = db.execute(q).fetchall()
    return [{"id": col[0], "name": col[1], "type": col[2]} for col in columns]

@app.get('/companies/all')
def list_companies(db: sqlite3.Connection = Depends(get_db)):
    q = "SELECT * FROM companies;"
    companies = db.execute(q).fetchall()
    return [dict(company) for company in companies]

@app.get('/tables/inspect/{table_name}')
def inspect_table(table_name: str, db: sqlite3.Connection = Depends(get_db)):
    q = f"PRAGMA table_info({table_name});"
    columns = db.execute(q).fetchall()
    return [{"id": col[0], "name": col[1], "type": col[2]} for col in columns]

# Endpoint untuk melihat daftar tabel dalam database
@app.get('/tables')
def list_tables():
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        return [table[0] for table in cursor.fetchall()]

